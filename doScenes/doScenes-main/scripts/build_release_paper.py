from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import ListFlowable, ListItem, PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle


@dataclass
class Heading:
    level: int
    text: str


@dataclass
class ParagraphBlock:
    text: str


@dataclass
class ListBlock:
    items: list[str]
    ordered: bool


@dataclass
class CodeBlock:
    text: str


@dataclass
class TableBlock:
    headers: list[str]
    rows: list[list[str]]


Block = Heading | ParagraphBlock | ListBlock | CodeBlock | TableBlock


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def normalize_inline(text: str) -> str:
    cleaned = LINK_RE.sub(r"\1 (\2)", text)
    cleaned = BOLD_RE.sub(r"\1", cleaned)
    cleaned = INLINE_CODE_RE.sub(r"\1", cleaned)
    return cleaned.strip()


def parse_markdown(markdown: str) -> list[Block]:
    lines = markdown.splitlines()
    blocks: list[Block] = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append(CodeBlock(text="\n".join(code_lines).rstrip()))
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            blocks.append(Heading(level=len(heading_match.group(1)), text=normalize_inline(heading_match.group(2))))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines):
            separator = lines[i + 1].strip()
            if separator.startswith("|") and "-" in separator:
                table_lines = [stripped]
                i += 1
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                headers = split_table_row(table_lines[0])
                rows = [split_table_row(row) for row in table_lines[2:]]
                blocks.append(TableBlock(headers=headers, rows=rows))
                continue

        if stripped.startswith("- "):
            items: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(normalize_inline(lines[i].strip()[2:]))
                i += 1
            blocks.append(ListBlock(items=items, ordered=False))
            continue

        if ordered_list_prefix(stripped):
            items = []
            while i < len(lines) and ordered_list_prefix(lines[i].strip()):
                item_text = re.sub(r"^\d+\.\s+", "", lines[i].strip(), count=1)
                items.append(normalize_inline(item_text))
                i += 1
            blocks.append(ListBlock(items=items, ordered=True))
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt.startswith("```") or HEADING_RE.match(nxt) or nxt.startswith("- ") or nxt.startswith("|") or ordered_list_prefix(nxt):
                break
            paragraph_lines.append(nxt)
            i += 1
        blocks.append(ParagraphBlock(text=normalize_inline(" ".join(paragraph_lines))))

    return blocks


def split_table_row(row: str) -> list[str]:
    cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
    return [normalize_inline(cell) for cell in cells]


def ordered_list_prefix(text: str) -> bool:
    return bool(re.match(r"^\d+\.\s+", text))


def set_run_font(run, font_name: str, size_pt: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def configure_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    style_map = {
        "Normal": ("Microsoft YaHei", 10.5, False, RGBColor(34, 34, 34)),
        "Title": ("Microsoft YaHei", 22, True, RGBColor(15, 51, 89)),
        "Subtitle": ("Microsoft YaHei", 12, False, RGBColor(92, 101, 112)),
        "Heading 1": ("Microsoft YaHei", 16, True, RGBColor(15, 51, 89)),
        "Heading 2": ("Microsoft YaHei", 13, True, RGBColor(26, 83, 118)),
        "Heading 3": ("Microsoft YaHei", 11.5, True, RGBColor(45, 62, 80)),
        "List Bullet": ("Microsoft YaHei", 10.5, False, RGBColor(34, 34, 34)),
        "List Number": ("Microsoft YaHei", 10.5, False, RGBColor(34, 34, 34)),
    }
    for style_name, (font_name, size_pt, bold, color) in style_map.items():
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        style.font.color.rgb = color


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_docx(blocks: list[Block], output_path: Path, author: str) -> None:
    doc = Document()
    configure_document_styles(doc)

    title = "Technical Report"
    if blocks and isinstance(blocks[0], Heading) and blocks[0].level == 1:
        title = blocks[0].text
        blocks = blocks[1:]

    title_p = doc.add_paragraph(style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    set_run_font(title_run, "Microsoft YaHei", 22, True, RGBColor(15, 51, 89))

    subtitle_p = doc.add_paragraph(style="Subtitle")
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("doScenes Challenge Language+History Track Technical Report")
    set_run_font(subtitle_run, "Microsoft YaHei", 11.5, False, RGBColor(92, 101, 112))

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run("Prepared from project experiments and submission artifacts")
    set_run_font(meta_run, "Microsoft YaHei", 10, False, RGBColor(92, 101, 112))
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_p.add_run(f"Author: {author}")
    set_run_font(author_run, "Microsoft YaHei", 10.5, True, RGBColor(45, 62, 80))

    doc.add_paragraph()
    summary_box = doc.add_table(rows=2, cols=3)
    summary_box.style = "Table Grid"
    summary_box.autofit = True
    labels = ["ADE_instruction (m)", "ADE_baseline (m)", "Delta ADE (m)"]
    values = ["3.474533", "3.955950", "+0.481417"]
    for idx, label in enumerate(labels):
        cell = summary_box.rows[0].cells[idx]
        shade_cell(cell, "DCE6F1")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, "Microsoft YaHei", 10.5, True, RGBColor(15, 51, 89))
    for idx, value in enumerate(values):
        cell = summary_box.rows[1].cells[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        set_run_font(r, "Microsoft YaHei", 11, True, RGBColor(34, 34, 34))

    doc.add_page_break()

    for block in blocks:
        if isinstance(block, Heading):
            style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(block.level, "Heading 3")
            p = doc.add_paragraph(style=style_name)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(block.text)
            if style_name == "Heading 1":
                set_run_font(run, "Microsoft YaHei", 16, True, RGBColor(15, 51, 89))
            elif style_name == "Heading 2":
                set_run_font(run, "Microsoft YaHei", 13, True, RGBColor(26, 83, 118))
            else:
                set_run_font(run, "Microsoft YaHei", 11.5, True, RGBColor(45, 62, 80))
        elif isinstance(block, ParagraphBlock):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.line_spacing = 1.35
            p.paragraph_format.space_after = Pt(5)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(block.text)
            set_run_font(run, "Microsoft YaHei", 10.5, False, RGBColor(34, 34, 34))
        elif isinstance(block, ListBlock):
            style_name = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                p = doc.add_paragraph(style=style_name)
                p.paragraph_format.line_spacing = 1.25
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(item)
                set_run_font(run, "Microsoft YaHei", 10.5, False, RGBColor(34, 34, 34))
        elif isinstance(block, CodeBlock):
            table = doc.add_table(rows=1, cols=1)
            table.autofit = True
            cell = table.cell(0, 0)
            shade_cell(cell, "F3F6F9")
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(block.text)
            set_run_font(run, "Consolas", 9.5, False, RGBColor(60, 60, 60))
        elif isinstance(block, TableBlock):
            table = doc.add_table(rows=1, cols=len(block.headers))
            table.style = "Table Grid"
            table.autofit = True
            header_cells = table.rows[0].cells
            for idx, header in enumerate(block.headers):
                shade_cell(header_cells[idx], "DCE6F1")
                p = header_cells[idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(header)
                set_run_font(r, "Microsoft YaHei", 10, True, RGBColor(15, 51, 89))
            for row in block.rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row):
                    p = cells[idx].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                    r = p.add_run(value)
                    set_run_font(r, "Microsoft YaHei", 10, False, RGBColor(34, 34, 34))
            doc.add_paragraph()

    doc.save(output_path)


def build_pdf(blocks: list[Block], output_path: Path, author: str) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BodyCN", parent=styles["BodyText"], fontName="STSong-Light", fontSize=10.5, leading=16, textColor=colors.HexColor("#222222"), alignment=TA_LEFT, spaceAfter=6))
    styles.add(ParagraphStyle(name="TitleCN", parent=styles["Title"], fontName="STSong-Light", fontSize=22, leading=28, alignment=TA_CENTER, textColor=colors.HexColor("#0F3359"), spaceAfter=10))
    styles.add(ParagraphStyle(name="SubTitleCN", parent=styles["BodyText"], fontName="STSong-Light", fontSize=11.5, leading=16, alignment=TA_CENTER, textColor=colors.HexColor("#5C6570"), spaceAfter=4))
    styles.add(ParagraphStyle(name="H1CN", parent=styles["Heading1"], fontName="STSong-Light", fontSize=16, leading=22, textColor=colors.HexColor("#0F3359"), spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name="H2CN", parent=styles["Heading2"], fontName="STSong-Light", fontSize=13, leading=18, textColor=colors.HexColor("#1A5376"), spaceBefore=8, spaceAfter=4))
    styles.add(ParagraphStyle(name="H3CN", parent=styles["Heading3"], fontName="STSong-Light", fontSize=11.5, leading=16, textColor=colors.HexColor("#2D3E50"), spaceBefore=6, spaceAfter=3))
    styles.add(ParagraphStyle(name="CodeCN", parent=styles["Code"], fontName="Courier", fontSize=8.5, leading=11, backColor=colors.HexColor("#F3F6F9"), leftIndent=8, rightIndent=8, borderPadding=6, borderColor=colors.HexColor("#D4DDE7"), borderWidth=0.5))

    story = []
    title = "Technical Report"
    if blocks and isinstance(blocks[0], Heading) and blocks[0].level == 1:
        title = blocks[0].text
        blocks = blocks[1:]

    story.append(Spacer(1, 2.2 * cm))
    story.append(Paragraph(escape(title), styles["TitleCN"]))
    story.append(Paragraph("doScenes Challenge Language+History Track Technical Report", styles["SubTitleCN"]))
    story.append(Paragraph("Prepared from project experiments and submission artifacts", styles["SubTitleCN"]))
    story.append(Paragraph(escape(f"Author: {author}"), styles["SubTitleCN"]))
    story.append(Spacer(1, 0.8 * cm))

    summary_table = Table(
        [
            ["ADE_instruction (m)", "ADE_baseline (m)", "Delta ADE (m)"],
            ["3.474533", "3.955950", "+0.481417"],
        ],
        colWidths=[4.2 * cm, 4.2 * cm, 4.2 * cm],
        hAlign="CENTER",
    )
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCE6F1")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0F3359")),
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("LEADING", (0, 0), (-1, -1), 14),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B4C6D6")),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(summary_table)
    story.append(PageBreak())

    for block in blocks:
        if isinstance(block, Heading):
            style = {1: styles["H1CN"], 2: styles["H2CN"], 3: styles["H3CN"]}.get(block.level, styles["H3CN"])
            story.append(Paragraph(escape(block.text), style))
        elif isinstance(block, ParagraphBlock):
            story.append(Paragraph(escape(block.text), styles["BodyCN"]))
        elif isinstance(block, ListBlock):
            bullet_style = styles["BodyCN"]
            flowable_items = [
                ListItem(Paragraph(escape(item), bullet_style), leftIndent=10)
                for item in block.items
            ]
            bullet_type = "1" if block.ordered else "bullet"
            story.append(ListFlowable(flowable_items, bulletType=bullet_type, start="1", leftIndent=16))
            story.append(Spacer(1, 0.15 * cm))
        elif isinstance(block, CodeBlock):
            story.append(Preformatted(block.text, styles["CodeCN"]))
            story.append(Spacer(1, 0.2 * cm))
        elif isinstance(block, TableBlock):
            data = [block.headers] + block.rows
            col_count = max(len(row) for row in data)
            usable_width = A4[0] - 4.6 * cm
            col_width = usable_width / col_count
            table = Table(data, colWidths=[col_width] * col_count, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCE6F1")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0F3359")),
                        ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("LEADING", (0, 0), (-1, -1), 12),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B4C6D6")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("ALIGN", (1, 1), (-1, -1), "CENTER"),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 0.2 * cm))

    doc = SimpleDocTemplate(str(output_path), pagesize=A4, leftMargin=2.4 * cm, rightMargin=2.2 * cm, topMargin=2.0 * cm, bottomMargin=2.0 * cm, title=title, author=author)
    doc.build(story)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build release paper DOCX and PDF from markdown.")
    parser.add_argument("--input", type=Path, required=True, help="Input markdown paper path.")
    parser.add_argument("--docx", type=Path, required=True, help="Output DOCX path.")
    parser.add_argument("--pdf", type=Path, required=True, help="Output PDF path.")
    parser.add_argument("--author", type=str, default="邓柯", help="Paper author name.")
    args = parser.parse_args()

    markdown = args.input.read_text(encoding="utf-8")
    blocks = parse_markdown(markdown)
    args.docx.parent.mkdir(parents=True, exist_ok=True)
    args.pdf.parent.mkdir(parents=True, exist_ok=True)
    build_docx(list(blocks), args.docx, args.author)
    build_pdf(list(blocks), args.pdf, args.author)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
